"""
Document Generator
Renders approved JBS JSON to a corporate Word (.docx) file,
uploads to Azure Blob Storage, and returns a SAS download URL.
"""

import os
import json
import uuid
from datetime import datetime, timezone, timedelta
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
from docx import Document
from docx.shared import Pt, RGBColor

TEMPLATE_PATH   = os.path.join(os.path.dirname(__file__), "../../templates/jbs_corporate_template.docx")
AZURE_ACCOUNT   = os.environ.get("AZURE_STORAGE_ACCOUNT", "")
AZURE_CONTAINER = os.environ.get("AZURE_STORAGE_CONTAINER", "certis-jbs-documents")
BLOB_PREFIX     = os.environ.get("BLOB_PREFIX", "jbs-documents/")
URL_EXPIRY      = int(os.environ.get("DOC_URL_EXPIRY_SECONDS", "900"))


class DocumentGenerator:
    def __init__(self):
        account_key = os.environ["AZURE_STORAGE_KEY"]
        conn_str = (
            f"DefaultEndpointsProtocol=https;"
            f"AccountName={AZURE_ACCOUNT};"
            f"AccountKey={account_key};"
            f"EndpointSuffix=core.windows.net"
        )
        self.blob_service = BlobServiceClient.from_connection_string(conn_str)
        self.account_key = account_key

    async def generate(self, jbs_json: dict) -> str:
        doc = Document(TEMPLATE_PATH)
        self._populate_document(doc, jbs_json)

        site = jbs_json["metadata"]["site_name"].replace(" ", "_")
        filename = f"JBS_{site}_{uuid.uuid4().hex[:8]}.docx"
        local_path = f"/tmp/{filename}"
        doc.save(local_path)

        blob_name = f"{BLOB_PREFIX}{filename}"
        container_client = self.blob_service.get_container_client(AZURE_CONTAINER)
        with open(local_path, "rb") as data:
            container_client.upload_blob(name=blob_name, data=data, overwrite=True)

        expiry = datetime.now(timezone.utc) + timedelta(seconds=URL_EXPIRY)
        sas_token = generate_blob_sas(
            account_name=AZURE_ACCOUNT,
            container_name=AZURE_CONTAINER,
            blob_name=blob_name,
            account_key=self.account_key,
            permission=BlobSasPermissions(read=True),
            expiry=expiry,
        )
        return (
            f"https://{AZURE_ACCOUNT}.blob.core.windows.net"
            f"/{AZURE_CONTAINER}/{blob_name}?{sas_token}"
        )

    @staticmethod
    def _all_paragraphs(doc: Document):
        """Yield every paragraph in the document, including those inside table cells."""
        yield from doc.paragraphs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    @staticmethod
    def _replace_in_paragraph(para, replacements: dict):
        """
        Replace placeholder text inside a paragraph while preserving run formatting.

        Word sometimes splits a placeholder across multiple runs (e.g. "{CUSTOMER" +
        "_NAME}"), so we operate on the full paragraph text to detect the placeholder,
        but write the replacement into individual runs to keep bold/colour/size intact.
        """
        if not any(k in para.text for k in replacements):
            return
        for run in para.runs:
            for placeholder, value in replacements.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

    def _populate_document(self, doc: Document, data: dict):
        meta = data["metadata"]
        replacements = {
            "{CUSTOMER_NAME}": meta.get("customer_name", ""),
            "{SITE_NAME}":     meta.get("site_name", ""),
            "{SITE_CATEGORY}": meta.get("site_category", ""),
            "{JOB_PURPOSE}":   meta.get("job_purpose", ""),
            "{GENERATED_AT}":  data.get("generated_at", ""),
            "{AUTHORIZED_BY}": meta.get("authorized_by", ""),
        }
        for para in self._all_paragraphs(doc):
            self._replace_in_paragraph(para, replacements)

        # Duties tables
        for duty in data.get("duties", []):
            doc.add_heading(duty["duty_name"], level=2)
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for i, h in enumerate(["#", "Task", "Trigger", "Frequency", "Role"]):
                hdr[i].text = h
            for task in duty.get("tasks", []):
                row = table.add_row().cells
                row[0].text = str(task.get("sequence", ""))
                row[1].text = task.get("task_description", "")
                row[2].text = task.get("trigger", "")
                row[3].text = task.get("frequency", "")
                row[4].text = task.get("responsible_role", "")

        # Safety section
        sc = data.get("safety_compliance", {})
        doc.add_heading("Safety & Compliance", level=1)
        doc.add_paragraph(f"Hazards: {', '.join(sc.get('site_hazards', []))}")
        doc.add_paragraph(f"PPE: {', '.join(sc.get('ppe_requirements', []))}")

        # Append machine-readable JSON as hidden audit trail
        doc.add_page_break()
        doc.add_heading("Appendix: Machine-Readable Data", level=1)
        p = doc.add_paragraph()
        run = p.add_run(f"<JBS_DATA>{json.dumps(data)}</JBS_DATA>")
        run.font.size = Pt(6)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
