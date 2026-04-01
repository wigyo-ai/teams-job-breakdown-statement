"""
Creates templates/jbs_corporate_template.docx
Run once: python templates/create_template.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        attrs = kwargs.get(edge, {})
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), attrs.get("val", "single"))
        tag.set(qn("w:sz"), str(attrs.get("sz", 4)))
        tag.set(qn("w:color"), attrs.get("color", "auto"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, size=11, color=None, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return run


doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles ────────────────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10)

# ── HEADER BAND ───────────────────────────────────────────────────────────────
header_table = doc.add_table(rows=1, cols=2)
header_table.style = "Table Grid"
header_table.autofit = False
header_table.columns[0].width = Inches(3.5)
header_table.columns[1].width = Inches(3.5)

left_cell = header_table.cell(0, 0)
right_cell = header_table.cell(0, 1)

set_cell_bg(left_cell, "1A3A5C")    # Certis navy
set_cell_bg(right_cell, "1A3A5C")

# Company name — left
lp = left_cell.paragraphs[0]
lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(lp, "CERTIS SECURITY", bold=True, size=14, color="FFFFFF")
lp.paragraph_format.space_before = Pt(8)
lp.paragraph_format.space_after = Pt(2)
add_run(left_cell.add_paragraph(), "Security Operations", size=9, color="A8C4E0")

# Document type — right
rp = right_cell.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(rp, "JOB BREAKDOWN STATEMENT", bold=True, size=13, color="FFFFFF")
rp.paragraph_format.space_before = Pt(8)
rp.paragraph_format.space_after = Pt(2)
rp2 = right_cell.add_paragraph()
rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(rp2, "AI-Assisted | Confidential", size=8, italic=True, color="A8C4E0")

doc.add_paragraph()  # spacer

# ── METADATA TABLE ────────────────────────────────────────────────────────────
meta_table = doc.add_table(rows=6, cols=2)
meta_table.style = "Table Grid"
meta_table.autofit = False
meta_table.columns[0].width = Inches(2.2)
meta_table.columns[1].width = Inches(4.8)

LABEL_COLOR  = "1A3A5C"
LABEL_TEXT   = "F0F4F8"
VALUE_COLOR  = "F8FAFC"

rows_data = [
    ("Customer / Client",  "{CUSTOMER_NAME}"),
    ("Site",               "{SITE_NAME}"),
    ("Site Category",      "{SITE_CATEGORY}"),
    ("Job Purpose",        "{JOB_PURPOSE}"),
    ("Generated",          "{GENERATED_AT}"),
    ("Authorized By",      "{AUTHORIZED_BY}"),
]

for i, (label, placeholder) in enumerate(rows_data):
    label_cell = meta_table.cell(i, 0)
    value_cell = meta_table.cell(i, 1)
    set_cell_bg(label_cell, LABEL_COLOR)
    set_cell_bg(value_cell, VALUE_COLOR)

    lp = label_cell.paragraphs[0]
    lp.paragraph_format.space_before = Pt(3)
    lp.paragraph_format.space_after  = Pt(3)
    add_run(lp, label, bold=True, size=9, color=LABEL_TEXT)

    vp = value_cell.paragraphs[0]
    vp.paragraph_format.space_before = Pt(3)
    vp.paragraph_format.space_after  = Pt(3)
    add_run(vp, placeholder, size=10)

doc.add_paragraph()  # spacer

# ── SECTION DIVIDER: DUTIES & TASKS ──────────────────────────────────────────
divider_table = doc.add_table(rows=1, cols=1)
divider_table.style = "Table Grid"
dc = divider_table.cell(0, 0)
set_cell_bg(dc, "E8F0F7")
dp = dc.paragraphs[0]
dp.alignment = WD_ALIGN_PARAGRAPH.LEFT
dp.paragraph_format.space_before = Pt(4)
dp.paragraph_format.space_after  = Pt(4)
add_run(dp, "DUTIES & TASKS", bold=True, size=11, color="1A3A5C")

intro = doc.add_paragraph()
add_run(
    intro,
    "The following duties and task sequences were captured through the AI-assisted "
    "interview process. Each duty is broken down into discrete tasks with their "
    "associated triggers, frequencies, and responsible roles.",
    size=9,
    italic=True,
    color="555555",
)
intro.paragraph_format.space_before = Pt(4)
intro.paragraph_format.space_after  = Pt(8)

# ── Save ──────────────────────────────────────────────────────────────────────
import os
out_path = os.path.join(os.path.dirname(__file__), "jbs_corporate_template.docx")
doc.save(out_path)
print(f"Template saved → {out_path}")
